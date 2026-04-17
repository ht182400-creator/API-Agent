#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""转换单个文档"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches

workspace = 'c:/Users/ht182/CodeBuddy/20260416205552'
md_file = os.path.join(workspace, '02_项目实施方案.md')
docx_file = os.path.join(workspace, '02_项目实施方案.docx')

# 删除旧的docx（如果存在）
if os.path.exists(docx_file):
    try:
        os.remove(docx_file)
    except:
        pass

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    if not stripped or stripped == '---':
        doc.add_paragraph()
        i += 1
        continue
    
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('#### '):
        doc.add_heading(stripped[5:], level=4)
    elif stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        data_lines = [ln for ln in table_lines if not re.match(r'^\|[\s\-:|]+\|$', ln)]
        if len(data_lines) >= 1 and data_lines[0]:
            try:
                header_cells = [c.strip() for c in data_lines[0].split('|')[1:-1]]
                if header_cells and len(data_lines) > 1:
                    table = doc.add_table(rows=len(data_lines)-1, cols=len(header_cells))
                    table.style = 'Table Grid'
                    for row_idx, table_line in enumerate(data_lines[1:]):
                        cells = [c.strip() for c in table_line.split('|')[1:-1]]
                        for j, cell_text in enumerate(cells):
                            if j < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[j].text = cell_text
            except: pass
        continue
    elif '```' in stripped:
        p = doc.add_paragraph()
        run = p.add_run(stripped.replace('`', ''))
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    else:
        p = doc.add_paragraph()
        p.add_run(stripped)
    i += 1

doc.save(docx_file)
print(f'Successfully saved: {docx_file}')
