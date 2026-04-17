#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MD文件转Word文档转换器
使用 python-docx 和 markdown 库
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import markdown

workspace = 'c:/Users/ht182/CodeBuddy/20260416205552'
md_files = [f for f in os.listdir(workspace) if f.endswith('.md')]

def create_docx(md_path, docx_path):
    """从MD文件创建Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # 读取MD文件
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 处理每一行
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 跳过空行和分隔线
        if not stripped or stripped == '---':
            if stripped:
                doc.add_paragraph()
            i += 1
            continue
        
        # 标题处理
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            i += 1
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            i += 1
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
            i += 1
        
        # 表格处理
        elif stripped.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 解析表格 - 过滤分隔行
            data_lines = [ln for ln in table_lines if not re.match(r'^\|[\s\-:|]+\|$', ln)]
            
            if len(data_lines) >= 1:
                try:
                    # 提取表头
                    header_cells = [c.strip() for c in data_lines[0].split('|')[1:-1]]
                    
                    if header_cells and len(data_lines) > 1:
                        # 创建表格
                        table = doc.add_table(rows=len(data_lines)-1, cols=len(header_cells))
                        table.style = 'Table Grid'
                        
                        # 填充数据
                        for row_idx, table_line in enumerate(data_lines[1:]):
                            cells = [c.strip() for c in table_line.split('|')[1:-1]]
                            for j, cell_text in enumerate(cells):
                                if j < len(table.rows[row_idx].cells):
                                    table.rows[row_idx].cells[j].text = cell_text
                except Exception as e:
                    # 如果表格解析失败，作为普通文本处理
                    for tline in data_lines:
                        p = doc.add_paragraph(tline)
            continue
        
        # 代码块
        elif '```' in stripped:
            code_lines = []
            i += 1
            while i < len(lines) and '```' not in lines[i]:
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # 有序列表
        elif re.match(r'^\d+\.', stripped):
            p = doc.add_paragraph(stripped, style='List Number')
            i += 1
        
        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            i += 1
        
        # 普通段落
        else:
            p = doc.add_paragraph()
            # 处理行内代码
            parts = re.split(r'`([^`]+)`', stripped)
            for part in parts:
                if re.match(r'^[\w\W]*`[\w\W]*`[\w\W]*$', stripped):
                    # 检查是否是代码部分
                    code_matches = re.findall(r'`([^`]+)`', stripped)
                    for code in code_matches:
                        stripped = stripped.replace(f'`{code}`', '', 1)
                        run = p.add_run(code)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    p.add_run(stripped)
                else:
                    p.add_run(part)
            i += 1
    
    # 保存文档
    doc.save(docx_path)

print(f'开始转换 {len(md_files)} 个MD文件为Word格式...\n')

success_count = 0
fail_count = 0

for md_file in md_files:
    md_path = os.path.join(workspace, md_file)
    docx_file = md_file.replace('.md', '.docx')
    docx_path = os.path.join(workspace, docx_file)
    
    try:
        create_docx(md_path, docx_path)
        print(f'[OK] {md_file} -> {docx_file}')
        success_count += 1
    except Exception as e:
        print(f'[FAIL] {md_file}: {str(e)}')
        fail_count += 1

print(f'\n转换完成: 成功 {success_count}, 失败 {fail_count}')
