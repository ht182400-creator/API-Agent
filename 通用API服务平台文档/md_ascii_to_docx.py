#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Markdown文件提取ASCII图表并转为Word标准表格
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_table_borders(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_header_style(cell):
    """设置表头样式"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E2F3')
    tcPr.append(shd)


def find_ascii_diagram_blocks(lines):
    """找出所有ASCII图表块"""
    blocks = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 检测ASCII图表开始行（包含┌┐└┘和─│字符）
        if '┌' in line and ('─' in line or '│' in line):
            block_start = i
            block_lines = []
            
            # 收集连续的相关行
            while i < len(lines):
                curr_line = lines[i]
                # 检查是否仍然是ASCII图表行
                if '┌' in curr_line or '┐' in curr_line or '└' in curr_line or '┘' in curr_line or \
                   ('│' in curr_line and '─' in curr_line) or \
                   re.match(r'^[\s│├┤┬┴┼─]+$', curr_line.strip()):
                    if curr_line.strip():  # 只添加非空行
                        block_lines.append(curr_line.rstrip())
                    i += 1
                elif curr_line.strip() == '':
                    # 空行可能是分隔
                    if block_lines:
                        block_lines.append('')
                    i += 1
                    # 如果连续空行，则结束当前块
                    if len(block_lines) > 0 and all(l.strip() == '' for l in block_lines[-3:]):
                        break
                else:
                    break
            
            if block_lines and any('┌' in l or '┐' in l or '└' in l or '┘' in l for l in block_lines):
                blocks.append({
                    'start': block_start,
                    'end': i,
                    'lines': block_lines
                })
            continue
        
        i += 1
    
    return blocks


def parse_ascii_diagram_to_table(lines):
    """将ASCII图表解析为表格数据"""
    # 提取每个框内的文字
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 匹配各种格式的单元格
        # 格式1: │ 内容 │ 内容 │
        # 格式2: │   内容   │
        row_cells = []
        
        # 分割行中的单元格
        parts = line.split('│')
        for part in parts:
            part = part.strip()
            if part and not all(c in '─┌┐└┘├┤┬┴┼' for c in part):
                # 清理内容
                content = part.replace('─', '').strip()
                if content:
                    row_cells.append(content)
        
        if row_cells:
            table_data.append(row_cells)
    
    return table_data


def process_markdown_file(md_file):
    """处理单个Markdown文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 找到所有ASCII图表块
    blocks = find_ascii_diagram_blocks(lines)
    
    if not blocks:
        print(f"  未发现ASCII图表")
        return 0
    
    print(f"  发现 {len(blocks)} 个ASCII图表")
    
    # 创建Word文档
    doc = Document()
    
    current_line = 0
    for block in blocks:
        # 添加图表前的内容
        for j in range(current_line, block['start']):
            line = lines[j].strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line and not line.startswith('```') and not line.startswith('|'):
                p = doc.add_paragraph()
                p.add_run(line)
        
        # 解析ASCII图表
        table_data = parse_ascii_diagram_to_table(block['lines'])
        
        if table_data:
            # 创建Word表格
            col_count = max(len(row) for row in table_data)
            table = doc.add_table(rows=len(table_data), cols=col_count)
            set_table_borders(table)
            
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        
                        if row_idx == 0:
                            set_header_style(cell)
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                        else:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(9)
            
            doc.add_paragraph()
        else:
            # 如果解析失败，直接添加原始ASCII（等宽字体）
            p = doc.add_paragraph()
            p.add_run('[ASCII图表]').italic = True
            for line in block['lines']:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
        
        current_line = block['end']
    
    # 添加剩余内容
    for j in range(current_line, len(lines)):
        line = lines[j].strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line and not line.startswith('```'):
            p = doc.add_paragraph()
            p.add_run(line)
    
    # 保存
    output_file = md_file.replace('.md', '_表格.docx')
    doc.save(output_file)
    print(f"  已保存: {os.path.basename(output_file)}")
    
    return len(blocks)


def main():
    """主函数"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(base_dir)
    
    files = [
        os.path.join(parent_dir, '18_通用API服务平台架构.md'),
        os.path.join(parent_dir, 'API服务改造项目需求分析报告.md'),
        os.path.join(base_dir, '19_通用API服务平台需求规格说明书.md'),
        os.path.join(base_dir, '20_通用API服务平台实施方案.md'),
        os.path.join(parent_dir, '01_项目需求规格说明书.md'),
        os.path.join(parent_dir, '02_项目实施方案.md'),
        os.path.join(parent_dir, '03_开发周期计划.md'),
    ]
    
    print("=" * 50)
    print("ASCII图表转Word标准表格")
    print("=" * 50)
    
    total = 0
    for f in files:
        if os.path.exists(f):
            print(f"\n处理: {os.path.basename(f)}")
            count = process_markdown_file(f)
            total += count
    
    print("\n" + "=" * 50)
    print(f"共转换 {total} 个图表")


if __name__ == '__main__':
    main()
