# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
import re
import os

def convert_md_to_docx(md_file):
    docx_file = md_file.replace('.md', '.docx')
    try:
        if os.path.exists(docx_file):
            os.remove(docx_file)
    except: pass
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped or stripped == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            data_lines = [ln for ln in table_lines if not re.match(r'^\|[\s\-\:\|\+]+\|$', ln)]
            if data_lines and len(data_lines) > 1:
                try:
                    headers = [c.strip() for c in data_lines[0].split('|')[1:-1]]
                    table = doc.add_table(rows=len(data_lines)-1, cols=len(headers))
                    table.style = 'Table Grid'
                    for row_idx, table_line in enumerate(data_lines[1:]):
                        cells = [c.strip() for c in table_line.split('|')[1:-1]]
                        for j, cell_text in enumerate(cells):
                            if j < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[j].text = cell_text
                except Exception as e:
                    print(f"Table error: {e}")
            continue
        elif stripped.startswith('```mermaid'):
            p = doc.add_paragraph()
            p.add_run('[Mermaid图表 - 请在Markdown预览中查看]')
        elif stripped.startswith('```') and not stripped.startswith('```mermaid'):
            continue
        else:
            p = doc.add_paragraph()
            p.add_run(stripped)
        i += 1
    
    doc.save(docx_file)
    print(f'OK: {os.path.basename(docx_file)}')

# Convert specific files
files = [
    '19_通用API服务平台需求规格说明书.md',
    '20_通用API服务平台实施方案.md',
    '26_通用API服务平台数据库设计文档.md',
    '28_通用API服务平台图表规范.md'
]

for f in files:
    if os.path.exists(f):
        convert_md_to_docx(f)
