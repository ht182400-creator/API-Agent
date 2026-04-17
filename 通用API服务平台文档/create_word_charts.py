#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ASCII图表转Word原生图表 - 简洁版
使用表格+样式模拟形状效果
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg_color(cell, color_hex):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_border_color(cell, color='000000'):
    """设置单元格边框颜色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    
    tcPr.append(tcBorders)


def add_styled_paragraph(doc, text, bg_color, text_color='FFFFFF', bold=True, font_size=10):
    """添加带样式的段落（模拟形状）"""
    p = doc.add_paragraph()
    
    # 添加带有背景色的文本
    run = p.add_run(f' {text} ')
    run.font.name = '微软雅黑'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(text_color)
    
    # 通过添加前后空格来模拟边距
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    return p


def create_flow_chart_table(doc, boxes, title='流程图'):
    """使用表格创建流程图"""
    doc.add_heading(title, level=2)
    
    # 创建2列表格：方框 + 箭头
    for i, box in enumerate(boxes[:6]):
        # 行表格
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽
        for cell in table.rows[0].cells:
            cell.width = Cm(4)
        
        # 第一个单元格：方框
        cell1 = table.rows[0].cells[0]
        set_cell_bg_color(cell1, '4472C4')
        set_cell_border_color(cell1, '2F5496')
        p = cell1.paragraphs[0]
        run = p.add_run(box.get('text', f'步骤{i+1}'))
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 第二个单元格：箭头
        cell2 = table.rows[0].cells[1]
        p = cell2.paragraphs[0]
        run = p.add_run(' → ')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 112, 192)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 第三个单元格：空白
        cell3 = table.rows[0].cells[2]
    
    doc.add_paragraph()


def create_layer_table(doc, title, layers):
    """使用表格创建分层架构图"""
    doc.add_heading(title, level=2)
    
    for name, desc, color in layers:
        # 创建带颜色的标题行
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell = table.rows[0].cells[0]
        set_cell_bg_color(cell, color)
        set_cell_border_color(cell, '000000')
        
        p = cell.paragraphs[0]
        run = p.add_run(f'{name}\n')
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph()


def create_comparison_table(doc, title='对比图'):
    """创建对比表格"""
    doc.add_heading(title, level=2)
    
    # 传统模式
    p = doc.add_paragraph()
    run = p.add_run('【传统模式】')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.size = Pt(11)
    
    # 添加对比行
    traditional = [
        ('开发者', '直接对接仓库A/B/C', 'C00000'),
        ('认证', '每个仓库单独认证', 'C00000'),
        ('计费', '无统一计费', 'C00000'),
        ('问题', '对接成本高、维护复杂', 'C00000'),
    ]
    
    table1 = doc.add_table(rows=len(traditional), cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (col1, col2, color) in enumerate(traditional):
        set_cell_bg_color(table1.rows[i].cells[0], 'FFE6E6')
        set_cell_border_color(table1.rows[i].cells[0], color)
        p = table1.rows[i].cells[0].paragraphs[0]
        run = p.add_run(col1)
        run.bold = True
        run.font.size = Pt(10)
        
        set_cell_bg_color(table1.rows[i].cells[1], 'FFFFFF')
        set_cell_border_color(table1.rows[i].cells[1], color)
        p = table1.rows[i].cells[1].paragraphs[0]
        run = p.add_run(col2)
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # 平台模式
    p = doc.add_paragraph()
    run = p.add_run('【平台模式 - 推荐】')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run.font.size = Pt(11)
    
    platform = [
        ('开发者', '一次对接，接入所有仓库', '70AD47'),
        ('认证', '统一认证/JWT/OAuth', '70AD47'),
        ('计费', '统一计费和配额管理', '70AD47'),
        ('优势', '降低对接成本、易于维护', '70AD47'),
    ]
    
    table2 = doc.add_table(rows=len(platform), cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (col1, col2, color) in enumerate(platform):
        set_cell_bg_color(table2.rows[i].cells[0], 'E2EFDA')
        set_cell_border_color(table2.rows[i].cells[0], color)
        p = table2.rows[i].cells[0].paragraphs[0]
        run = p.add_run(col1)
        run.bold = True
        run.font.size = Pt(10)
        
        set_cell_bg_color(table2.rows[i].cells[1], 'FFFFFF')
        set_cell_border_color(table2.rows[i].cells[1], color)
        p = table2.rows[i].cells[1].paragraphs[0]
        run = p.add_run(col2)
        run.font.size = Pt(10)
    
    doc.add_paragraph()


def create_timeline_table(doc, phases, title='时间线'):
    """创建时间线表格"""
    doc.add_heading(title, level=2)
    
    table = doc.add_table(rows=2, cols=len(phases))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    colors = ['4472C4', '70AD47', 'ED7D31', '7030A0', 'C00000']
    
    # 第一行：阶段名
    for i, (phase, duration) in enumerate(phases):
        cell = table.rows[0].cells[i]
        set_cell_bg_color(cell, colors[i % len(colors)])
        set_cell_border_color(cell, colors[i % len(colors)])
        p = cell.paragraphs[0]
        run = p.add_run(phase)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 第二行：工期
    for i, (phase, duration) in enumerate(phases):
        cell = table.rows[1].cells[i]
        set_cell_bg_color(cell, 'F2F2F2')
        p = cell.paragraphs[0]
        run = p.add_run(duration)
        run.font.size = Pt(9)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()


def detect_diagram_type(block):
    """检测图表类型"""
    combined = '\n'.join(block)
    
    if '传统模式' in combined and '平台模式' in combined:
        return 'comparison'
    elif '接入层' in combined or 'Gateway' in combined or '架构' in combined:
        return 'architecture'
    elif '──▶' in combined or '调用流程' in combined:
        return 'flowchart'
    elif '开发' in combined and ('周期' in combined or '工期' in combined):
        return 'timeline'
    else:
        return 'architecture'


def process_file(md_file):
    """处理单个文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 查找图表块
    diagram_blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if '┌' in line and ('─' in line or '│' in line):
            block = []
            j = i
            while j < len(lines):
                l = lines[j]
                if '┌' in l or '┐' in l or '└' in l or '┘' in l or '│' in l or '─' in l or l.strip() == '':
                    if l.strip():
                        block.append(l)
                    j += 1
                else:
                    break
            if len(block) >= 2:
                diagram_blocks.append((i, j, block))
            i = j
        else:
            i += 1
    
    if not diagram_blocks:
        print(f"  未发现ASCII图表")
        return False
    
    # 创建文档
    doc = Document()
    doc.add_heading(f'ASCII图表可视化 - {os.path.basename(md_file)}', level=1)
    
    print(f"  发现 {len(diagram_blocks)} 个图表")
    
    for idx, (start, end, block) in enumerate(diagram_blocks):
        diagram_type = detect_diagram_type(block)
        
        # 添加原ASCII版本
        doc.add_heading(f'图表 {idx + 1} - {diagram_type}', level=2)
        
        p = doc.add_paragraph()
        run = p.add_run('[原始ASCII版本]')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        for line in block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph()
        
        # 创建可视化版本
        p = doc.add_paragraph()
        run = p.add_run('[可视化版本]')
        run.italic = True
        run.font.color.rgb = RGBColor(0, 128, 0)
        
        if diagram_type == 'comparison':
            create_comparison_table(doc, f'对比图 {idx + 1}')
        elif diagram_type == 'architecture':
            layers = [
                ('接入层', 'HTTPS | WAF | 限流 | 路由', '4472C4'),
                ('认证层', 'API Key | JWT | OAuth', '70AD47'),
                ('核心服务', '路由引擎 | 协议转换 | 计费', 'ED7D31'),
                ('适配层', 'HTTP | WebSocket | gRPC', '7030A0'),
                ('数据层', 'PostgreSQL | Redis', 'C00000'),
            ]
            create_layer_table(doc, f'架构图 {idx + 1}', layers)
        elif diagram_type == 'flowchart':
            boxes = []
            for line in block:
                parts = re.findall(r'│([^│]+)│', line)
                for part in parts:
                    text = part.strip()
                    if text and len(text) < 15:
                        boxes.append({'text': text})
            if not boxes:
                boxes = [{'text': f'步骤{i+1}'} for i in range(5)]
            create_flow_chart_table(doc, boxes, f'流程图 {idx + 1}')
        elif diagram_type == 'timeline':
            phases = [
                ('阶段一\n项目初始化', '3天'),
                ('阶段二\n核心功能', '2周'),
                ('阶段三\n扩展功能', '2周'),
                ('阶段四\n商业化', '1周'),
                ('阶段五\n上线运维', '持续'),
            ]
            create_timeline_table(doc, phases, f'时间线 {idx + 1}')
        else:
            create_layer_table(doc, f'架构图 {idx + 1}', layers if 'layers' in dir() else [])
        
        doc.add_page_break()
    
    # 保存
    output = md_file.replace('.md', '_可视化图表.docx')
    doc.save(output)
    print(f"  已保存: {os.path.basename(output)}")
    
    return True


def main():
    """主函数"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(base_dir)
    
    docs = [
        os.path.join(parent_dir, '18_通用API服务平台架构.md'),
        os.path.join(parent_dir, 'API服务改造项目需求分析报告.md'),
        os.path.join(base_dir, '19_通用API服务平台需求规格说明书.md'),
        os.path.join(base_dir, '20_通用API服务平台实施方案.md'),
        os.path.join(parent_dir, '01_项目需求规格说明书.md'),
        os.path.join(parent_dir, '02_项目实施方案.md'),
        os.path.join(parent_dir, '03_开发周期计划.md'),
    ]
    
    print("=" * 60)
    print("ASCII图表转Word原生图表")
    print("=" * 60)
    
    count = 0
    for doc_path in docs:
        if os.path.exists(doc_path):
            print(f"\n处理: {os.path.basename(doc_path)}")
            if process_file(doc_path):
                count += 1
    
    print("\n" + "=" * 60)
    print(f"完成! 成功处理 {count} 个文档")
    print("=" * 60)


if __name__ == '__main__':
    main()
