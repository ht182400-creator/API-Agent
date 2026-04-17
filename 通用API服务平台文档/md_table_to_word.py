#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown表格转Word原生表格工具
将文档中的|表格|格式转换为Word原生表格格式
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)


def is_markdown_table_line(line):
    """判断是否为Markdown表格行"""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    # 排除分隔行（包含---）
    if re.match(r'^\|[\s\-:\+|]+\|$', stripped):
        return False
    return True


def is_separator_line(line):
    """判断是否为Markdown表格分隔行"""
    stripped = line.strip()
    return bool(re.match(r'^\|[\s\-:\+|]+\|$', stripped))


def parse_markdown_table(table_lines):
    """解析Markdown表格内容"""
    rows = []
    for line in table_lines:
        if is_separator_line(line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]  # 去掉首尾的空字符串
        rows.append(cells)
    return rows


def create_word_table(doc, rows, header_color='D9E2F3'):
    """创建Word原生表格"""
    if not rows:
        return None
    
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    
    # 设置表格宽度
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text
                
                # 表头样式
                if row_idx == 0:
                    set_cell_shading(cell, header_color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                
                # 常规单元格样式
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                
                # 设置单元格对齐
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    return table


def convert_markdown_to_word_tables(md_file):
    """将Markdown文件中的表格转换为Word原生表格"""
    print(f"\n处理文件: {md_file}")
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 找到所有表格的位置
    table_ranges = []
    i = 0
    while i < len(lines):
        if is_markdown_table_line(lines[i]):
            table_start = i
            table_lines = []
            while i < len(lines) and (is_markdown_table_line(lines[i]) or is_separator_line(lines[i])):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:  # 至少2行（表头+数据）
                table_ranges.append((table_start, i, table_lines))
            continue
        i += 1
    
    if not table_ranges:
        print(f"  未找到表格")
        return content
    
    print(f"  找到 {len(table_ranges)} 个表格")
    
    # 创建新内容，将表格行替换为标记
    new_lines = []
    last_end = 0
    for start, end, table_lines in table_ranges:
        new_lines.extend(lines[last_end:start])
        new_lines.append(f'<!-- TABLE_START_{len(table_ranges.index((start,end,table_lines)))} -->')
        last_end = end
    
    new_lines.extend(lines[last_end:])
    
    return '\n'.join(new_lines), table_ranges


def process_all_documents():
    """处理所有文档"""
    doc_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Markdown文件列表
    md_files = [
        os.path.join(doc_dir, '19_通用API服务平台需求规格说明书.md'),
        os.path.join(doc_dir, '20_通用API服务平台实施方案.md'),
        os.path.join(doc_dir, '26_通用API服务平台数据库设计文档.md'),
    ]
    
    for md_file in md_files:
        if not os.path.exists(md_file):
            print(f"文件不存在: {md_file}")
            continue
        
        print(f"\n{'='*60}")
        print(f"处理文件: {os.path.basename(md_file)}")
        
        # 读取文件
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        # 找到所有表格
        table_ranges = []
        i = 0
        while i < len(lines):
            if is_markdown_table_line(lines[i]):
                table_start = i
                table_lines = []
                while i < len(lines) and (is_markdown_table_line(lines[i]) or is_separator_line(lines[i])):
                    table_lines.append(lines[i])
                    i += 1
                if len(table_lines) >= 2:
                    table_ranges.append((table_start, i, table_lines))
                continue
            i += 1
        
        if not table_ranges:
            print(f"  未找到表格")
            continue
        
        print(f"  找到 {len(table_ranges)} 个表格")
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        
        current_line = 0
        table_index = 0
        
        for start, end, table_lines in table_ranges:
            # 添加表格前的文本
            for j in range(current_line, start):
                line = lines[j].strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('```mermaid'):
                    p = doc.add_paragraph()
                    p.add_run('[Mermaid图表]').italic = True
                else:
                    p = doc.add_paragraph()
                    p.add_run(line)
            
            # 添加表格
            rows = parse_markdown_table(table_lines)
            if rows:
                create_word_table(doc, rows)
                print(f"    转换表格 {table_index + 1}: {len(rows)}行 x {len(rows[0]) if rows else 0}列")
            
            current_line = end
            table_index += 1
        
        # 添加剩余内容
        for j in range(current_line, len(lines)):
            line = lines[j].strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif '```mermaid' in line:
                p = doc.add_paragraph()
                p.add_run('[Mermaid图表]').italic = True
            else:
                p = doc.add_paragraph()
                p.add_run(line)
        
        # 保存文档
        docx_file = md_file.replace('.md', '_with_tables.docx')
        doc.save(docx_file)
        print(f"  已保存: {os.path.basename(docx_file)}")


def create_sample_table_document():
    """创建一个包含所有转换表格的示例文档"""
    doc_dir = os.path.dirname(os.path.abspath(__file__))
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    doc.add_heading('通用API服务平台 - 数据表格汇总', level=1)
    
    # 定义表格数据
    tables_data = [
        # 表1: 文档信息
        ('文档属性', [
            ('属性', '内容'),
            ('文档编号', 'REQ-PLATFORM-2026-001'),
            ('版本', 'V1.1'),
            ('日期', '2026-04-16'),
            ('项目名称', '通用API服务平台'),
        ]),
        # 表2: 术语定义
        ('核心术语定义', [
            ('术语', '定义'),
            ('平台', '通用API服务平台，作为API聚合中转站'),
            ('仓库', '可接入平台的API服务单元'),
            ('内部仓库', '平台自建的API服务'),
            ('外部仓库', '第三方接入的API服务'),
            ('开发者', '平台API的使用者'),
            ('仓库所有者', '外部仓库的管理者/提供者'),
        ]),
        # 表3: 痛点分析
        ('用户痛点分析', [
            ('痛点', '描述'),
            ('碎片化', '开发者需要在多个平台注册、管理多个API Key'),
            ('标准不统一', '各API服务接口规范不一致，集成成本高'),
            ('认证复杂', '每个平台都有自己的认证方式，管理繁琐'),
            ('计费混乱', '缺乏统一的计费和配额管理'),
        ]),
        # 表4: 仓库类型
        ('仓库类型说明', [
            ('类型', '说明', '示例'),
            ('内部仓库', '平台自建的服务', '课程问答、翻译服务'),
            ('外部仓库', '第三方接入的服务', '心理问答、股票查询'),
        ]),
        # 表5: 性能指标
        ('性能指标要求', [
            ('指标', '要求', '说明'),
            ('API响应时间', 'P99 < 500ms', '端到端响应时间'),
            ('吞吐量', '支持 10,000 QPS', '峰值处理能力'),
            ('并发连接', '支持 5,000 并发', 'WebSocket支持'),
            ('可用性', '99.9%', '年度停机时间 < 8.7小时'),
        ]),
        # 表6: 安全需求
        ('安全需求', [
            ('需求', '要求'),
            ('传输安全', 'HTTPS全链路加密'),
            ('认证安全', 'API Key + HMAC / JWT'),
            ('数据安全', '敏感数据加密存储'),
            ('审计日志', '完整操作日志记录'),
            ('防攻击', 'DDoS防护、限流熔断'),
        ]),
    ]
    
    # 添加表格
    for title, rows in tables_data:
        doc.add_heading(title, level=2)
        create_word_table(doc, rows)
        doc.add_paragraph()
    
    # 保存文档
    output_file = os.path.join(doc_dir, '表格汇总_原生Word格式.docx')
    doc.save(output_file)
    print(f"\n已创建表格汇总文档: {output_file}")


if __name__ == '__main__':
    print("=" * 60)
    print("Markdown表格转Word原生表格工具")
    print("=" * 60)
    
    # 处理所有文档
    process_all_documents()
    
    # 创建示例文档
    create_sample_table_document()
    
    print("\n" + "=" * 60)
    print("转换完成！")
    print("=" * 60)
