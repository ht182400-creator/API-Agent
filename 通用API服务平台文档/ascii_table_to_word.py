#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ASCII表格转Word标准表格 - 简洁版
直接将ASCII字符绘制的表格转换为Word原生表格
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def is_ascii_table_line(line):
    """判断是否是ASCII表格行"""
    return bool(re.search(r'[│├┤┬┴┼└┘┐]', line)) and bool(re.search(r'[─│]', line))


def is_separator_line(line):
    """判断是否是分隔行"""
    return bool(re.match(r'^[\s│├┤┬┴┼└┘┐]+$', line)) and '─' in line


def parse_ascii_table(lines):
    """解析ASCII表格"""
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if is_ascii_table_line(stripped):
            table_lines.append(stripped)
        elif is_separator_line(stripped):
            continue  # 跳过分隔行
        else:
            break
    
    return table_lines


def extract_cells(line):
    """从表格行中提取单元格内容"""
    cells = []
    
    # 匹配 │ 内容 │ 格式
    pattern = r'│([^│]*)│'
    matches = re.findall(pattern, line)
    
    for match in matches:
        cell_text = match.strip()
        # 清理ASCII边框字符
        cell_text = cell_text.replace('─', '').replace('│', '').strip()
        cells.append(cell_text)
    
    return cells


def set_table_borders(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_header_style(cell):
    """设置表头样式"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E2F3')
    tcPr.append(shd)


def convert_docx_ascii_tables(input_file, output_file=None):
    """转换单个Word文档中的ASCII表格"""
    if output_file is None:
        output_file = input_file
    
    # 读取源文档
    doc = Document(input_file)
    
    # 处理每个段落
    new_paragraphs = []
    skip_until = None
    
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text
        
        # 检查是否是ASCII表格开始
        if is_ascii_table_line(text):
            # 收集连续的表格行
            table_lines = []
            j = i
            while j < len(doc.paragraphs):
                p = doc.paragraphs[j]
                t = p.text.strip()
                if is_ascii_table_line(t) and not is_separator_line(t):
                    table_lines.append(t)
                    j += 1
                elif t == '' and table_lines:
                    # 空行可能是表格结束
                    if j + 1 < len(doc.paragraphs):
                        next_text = doc.paragraphs[j + 1].text.strip()
                        if not is_ascii_table_line(next_text):
                            break
                    j += 1
                else:
                    break
            
            if len(table_lines) >= 2:
                # 解析表格数据
                rows_data = []
                for line in table_lines:
                    cells = extract_cells(line)
                    if cells:
                        rows_data.append(cells)
                
                if rows_data:
                    # 创建Word表格
                    col_count = max(len(row) for row in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=col_count)
                    set_table_borders(table)
                    
                    for row_idx, row_data in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                
                                # 表头样式
                                if row_idx == 0:
                                    set_header_style(cell)
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.bold = True
                                            run.font.size = Pt(10)
                                else:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            run.font.size = Pt(9)
                    
                    i = j
                    continue
        else:
            new_paragraphs.append(para)
        
        i += 1
    
    # 保存
    doc.save(output_file)
    return True


def process_all_word_files():
    """处理所有Word文件"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 查找所有docx文件
    docx_files = [f for f in os.listdir(base_dir) if f.endswith('.docx') and not f.startswith('~$')]
    
    print("=" * 50)
    print("ASCII表格转Word标准表格")
    print("=" * 50)
    
    for filename in docx_files:
        filepath = os.path.join(base_dir, filename)
        print(f"\n处理: {filename}")
        
        try:
            convert_docx_ascii_tables(filepath)
            print(f"  完成")
        except Exception as e:
            print(f"  错误: {e}")
    
    print("\n" + "=" * 50)
    print("完成!")


if __name__ == '__main__':
    process_all_word_files()
