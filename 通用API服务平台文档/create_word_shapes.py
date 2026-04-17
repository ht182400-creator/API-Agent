#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ASCII图表转Word原生图表工具 - 使用底层XML API
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


# 形状类型常量
MSO_SHAPE_TYPE = {
    'ROUNDED_RECTANGLE': 5,
    'RECTANGLE': 1,
    'DIAMOND': 4,
    'OVAL': 9,
    'PENTAGON': 12,
    'HEXAGON': 10,
    'FLOWCHART_PROCESS': 11,
    'FLOWCHART_DECISION': 2,
    'FLOWCHART_DATA': 13,
    'FLOWCHART_DOCUMENT': 14,
    'FLOWCHART_TERMINATOR': 15,
    'CHART': 3,
    'NOTE': 25,
    'CALLOUT': 32,
    'ACTION_BUTTON_BACK': 37,
    'ACTION_BUTTON_NEXT': 38,
    'ACTION_BUTTON_END': 39,
}


def add_rounded_rectangle(doc, left, top, width, height, text, fill_color='4472C4'):
    """添加圆角矩形"""
    shape_id = doc.inline_shapes._get_next_id()
    
    # 创建图形元素
    nsmap_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    nsmap_p = 'http://schemas.openxmlformats.org/presentationml/2006/main'
    nsmap_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    
    # 创建shape XML
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)
    
    # 创建drawing元素
    drawing = OxmlElement('w:drawing')
    inline = OxmlElement('wp:inline')
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '0')
    inline.set('distR', '0')
    
    # 创建graphic
    graphic = OxmlElement('a:graphic')
    graphic_data = OxmlElement('a:graphicData')
    graphic_data.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/shape')
    
    # 创建形状
    shape = OxmlElement('p:sp')
    
    # nvSpPr
    nvSpPr = OxmlElement('p:nvSpPr')
    cNvPr = OxmlElement('p:cNvPr')
    cNvPr.set('id', str(shape_id))
    cNvPr.set('name', f'Rounded Rectangle {shape_id}')
    nvSpPr.append(cNvPr)
    
    cNvSpPr = OxmlElement('p:cNvSpPr')
    nvSpPr.append(cNvSpPr)
    shape.append(nvSpPr)
    
    # spPr
    spPr = OxmlElement('p:spPr')
    
    # 定义形状
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', str(int(left)))
    off.set('y', str(int(top)))
    xfrm.append(off)
    
    ext = OxmlElement('a:ext')
    ext.set('cx', str(int(width)))
    ext.set('cy', str(int(height)))
    xfrm.append(ext)
    spPr.append(xfrm)
    
    # 预设几何图形
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'roundRect')
    avLst = OxmlElement('a:avLst')
    prstGeom.append(avLst)
    spPr.append(prstGeom)
    
    # 填充颜色
    solidFill = OxmlElement('a:solidFill')
    srgbClr = OxmlElement('a:srgbClr')
    srgbClr.set('val', fill_color)
    solidFill.append(srgbClr)
    spPr.append(solidFill)
    
    # 线条
    ln = OxmlElement('a:ln')
    ln.append(solidFill)
    spPr.append(ln)
    
    shape.append(spPr)
    
    # txBody
    txBody = OxmlElement('p:txBody')
    bodyPr = OxmlElement('a:bodyPr')
    bodyPr.set('wrap', 'square')
    bodyPr.set('anchor', 'ctr')
    txBody.append(bodyPr)
    
    lstStyle = OxmlElement('a:lstStyle')
    txBody.append(lstStyle)
    
    para = OxmlElement('a:p')
    pPr = OxmlElement('a:pPr')
    pPr.set('algn', 'ctr')
    para.append(pPr)
    
    # 文本
    for line in text.split('\n'):
        r = OxmlElement('a:r')
        rPr = OxmlElement('a:rPr')
        rPr.set('lang', 'zh-CN')
        rPr.set('sz', '1200')
        rPr.set('b', '1')
        rPr.append(OxmlElement('a:solidFill'))
        rPr[-1].append(OxmlElement('a:srgbClr')).set('val', 'FFFFFF')
        r.append(rPr)
        
        t = OxmlElement('a:t')
        t.text = line
        r.append(t)
        
        para.append(r)
        
        # 换行
        if line != text.split('\n')[-1]:
            endParaRPr = OxmlElement('a:endParaRPr')
            para.append(endParaRPr)
    
    txBody.append(para)
    shape.append(txBody)
    
    graphic_data.append(shape)
    graphic.append(graphic_data)
    
    inline.append(graphic)
    drawing.append(inline)
    p.append(drawing)
    
    doc._body._element.append(p)
    
    return p


def add_simple_text_box(doc, text, left, top, width, height, fill_color='D9D9D9'):
    """添加简单文本框（带背景色）"""
    p = doc.add_paragraph()
    
    # 设置段落后退并添加表格来实现背景色
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    # 使用run来显示文本
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 添加边框效果（通过添加特殊字符）
    return p


def create_flowchart_diagram(doc, boxes, title='流程图'):
    """创建流程图"""
    doc.add_heading(title, level=2)
    
    # 布局参数
    box_width = Cm(4)
    box_height = Cm(1.5)
    h_spacing = Cm(1)
    v_spacing = Cm(0.5)
    
    # 分行显示
    rows = []
    current_row = []
    current_width = 0
    
    for i, box in enumerate(boxes[:8]):  # 限制最多8个
        if current_width + box_width > Cm(20):
            rows.append(current_row)
            current_row = [box]
            current_width = box_width
        else:
            current_row.append(box)
            current_width += box_width + h_spacing
    
    if current_row:
        rows.append(current_row)
    
    # 绘制
    start_top = Cm(5)
    
    for row_idx, row in enumerate(rows):
        start_left = Cm(2) + (Cm(20) - len(row) * box_width - (len(row) - 1) * h_spacing) / 2
        
        for col_idx, box in enumerate(row):
            left = start_left + col_idx * (box_width + h_spacing)
            top = start_top + row_idx * (box_height + v_spacing)
            
            # 添加方框
            add_rounded_rectangle(doc, left, top, box_width, box_height, 
                                  box.get('text', f'步骤{row_idx * len(row) + col_idx + 1}'),
                                  box.get('color', '4472C4'))
            
            # 添加箭头（如果不是最后一个）
            if col_idx < len(row) - 1:
                arrow_left = left + box_width
                arrow_top = top + box_height / 2
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = arrow_left
                run = p.add_run('→')
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 112, 192)
        
        # 行间箭头
        if row_idx < len(rows) - 1:
            down_top = start_top + row_idx * (box_height + v_spacing) + box_height
            p = doc.add_paragraph()
            run = p.add_run('↓')
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 112, 192)
            p.paragraph_format.left_indent = Cm(10)
    
    doc.add_paragraph()


def create_comparison_diagram(doc, title='对比图'):
    """创建对比图"""
    doc.add_heading(title, level=2)
    
    # 传统模式
    p = doc.add_paragraph()
    run = p.add_run('【传统模式】')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    # 添加开发者
    for i, name in enumerate(['开发者A', '开发者B', '开发者C']):
        p = doc.add_paragraph()
        run = p.add_run(f'  ┌──────────┐')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    for i, name in enumerate(['仓库A', '仓库B', '仓库C']):
        p = doc.add_paragraph()
        run = p.add_run(f'  └──────────┘')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 平台模式
    p = doc.add_paragraph()
    run = p.add_run('【平台模式】')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    # 添加图形化的平台框
    add_rounded_rectangle(doc, Cm(5), Cm(10), Cm(6), Cm(3),
                          'API聚合平台\n统一认证 | 路由分发\n计费统计 | 监控告警',
                          '4472C4')
    
    # 开发者
    for i, name in enumerate(['开发者A', '开发者B', '开发者C']):
        add_rounded_rectangle(doc, Cm(2), Cm(10.5) + i * Cm(1.2), Cm(2.5), Cm(0.9), name, '70AD47')
    
    # 仓库
    for i, name in enumerate(['仓库A', '仓库B', '仓库C']):
        add_rounded_rectangle(doc, Cm(12), Cm(10.5) + i * Cm(1.2), Cm(2.5), Cm(0.9), name, 'ED7D31')
    
    doc.add_paragraph()


def create_architecture_diagram(doc, title='架构图'):
    """创建架构图"""
    doc.add_heading(title, level=2)
    
    # 层
    layers = [
        ('接入层', 'HTTPS | WAF | 限流 | 路由', '4472C4'),
        ('认证层', 'API Key | JWT | OAuth', '70AD47'),
        ('核心服务', '路由引擎 | 协议转换 | 计费引擎', 'ED7D31'),
        ('适配层', 'HTTP适配器 | WebSocket | gRPC', '7030A0'),
        ('数据层', 'PostgreSQL | Redis | ClickHouse', 'C00000'),
    ]
    
    box_width = Cm(16)
    box_height = Cm(1.8)
    start_left = Cm(3)
    start_top = Cm(5)
    spacing = Cm(0.3)
    
    for i, (name, desc, color) in enumerate(layers):
        top = start_top + i * (box_height + spacing)
        text = f'{name}\n{desc}'
        add_rounded_rectangle(doc, start_left, top, box_width, box_height, text, color)
    
    # 添加右侧说明
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(20)
    run = p.add_run('┌─ 架构特点 ─────────────────┐')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    features = ['│ • 微服务架构，支持水平扩展', '│ • 插件化设计，热插拔仓库', '│ • 统一认证和计费', '│ • 高可用，支持多区域部署']
    for feature in features:
        p = doc.add_paragraph()
        run = p.add_run(feature)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run('└─────────────────────────────┘')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_paragraph()


def process_file_with_shapes(md_file):
    """处理单个文件，生成带形状的Word文档"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 查找图表块
    diagram_blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.search(r'┌', line) and ('─' in line or '│' in line):
            block = []
            j = i
            while j < len(lines):
                l = lines[j]
                if re.search(r'[┌┐└┘│├┤┬┴┼─]', l) or l.strip() == '':
                    block.append(l)
                    j += 1
                else:
                    break
            if len(block) >= 2:
                diagram_blocks.append((i, j, block))
            i = j
        else:
            i += 1
    
    if not diagram_blocks:
        print(f"  未发现ASCII图表")
        return False
    
    # 创建文档
    doc = Document()
    
    # 标题
    doc.add_heading(f'图表可视化文档', level=1)
    p = doc.add_paragraph()
    run = p.add_run(f'源文件: {os.path.basename(md_file)}')
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    print(f"  发现 {len(diagram_blocks)} 个图表")
    
    for idx, (start, end, block) in enumerate(diagram_blocks):
        doc.add_page_break()
        doc.add_heading(f'图表 {idx + 1}', level=2)
        
        # 显示原始ASCII
        p = doc.add_paragraph()
        run = p.add_run('【原始ASCII版本】')
        run.bold = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        for line in block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            p.paragraph_format.space_after = Pt(0)
        
        doc.add_paragraph()
        
        # 根据内容创建对应的可视化图表
        combined = '\n'.join(block)
        
        p = doc.add_paragraph()
        run = p.add_run('【可视化版本】')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)
        
        if '传统模式' in combined or '平台模式' in combined:
            create_comparison_diagram(doc, f'对比图 {idx + 1}')
        elif '接入层' in combined or 'Gateway' in combined:
            create_architecture_diagram(doc, f'架构图 {idx + 1}')
        elif '──▶' in combined or '调用' in combined:
            # 提取步骤
            boxes = []
            for line in block:
                parts = re.findall(r'│([^│]+)│', line)
                for part in parts:
                    text = part.strip()
                    if text and len(text) < 20:
                        boxes.append({'text': text})
            create_flowchart_diagram(doc, boxes, f'流程图 {idx + 1}')
        else:
            create_architecture_diagram(doc, f'架构图 {idx + 1}')
    
    # 保存
    output = md_file.replace('.md', '_可视化图表.docx')
    doc.save(output)
    print(f"  已保存: {os.path.basename(output)}")
    
    return True


def main():
    """主函数"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(base_dir)
    
    docs = [
        os.path.join(parent_dir, '18_通用API服务平台架构.md'),
        os.path.join(parent_dir, 'API服务改造项目需求分析报告.md'),
        os.path.join(base_dir, '19_通用API服务平台需求规格说明书.md'),
        os.path.join(base_dir, '20_通用API服务平台实施方案.md'),
        os.path.join(parent_dir, '01_项目需求规格说明书.md'),
        os.path.join(parent_dir, '02_项目实施方案.md'),
        os.path.join(parent_dir, '03_开发周期计划.md'),
    ]
    
    print("=" * 60)
    print("ASCII图表转Word原生图表 - 可视化版本")
    print("=" * 60)
    
    count = 0
    for doc_path in docs:
        if os.path.exists(doc_path):
            print(f"\n处理: {doc_path}")
            if process_file_with_shapes(doc_path):
                count += 1
    
    print("\n" + "=" * 60)
    print(f"完成! 成功处理 {count} 个文档")
    print("=" * 60)


if __name__ == '__main__':
    main()
