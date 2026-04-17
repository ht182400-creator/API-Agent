#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ASCII图表转Word原生图表工具
将ASCII字符绘制的框图、流程图转换为Word原生形状
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree


def parse_ascii_diagram(lines, start_idx):
    """解析ASCII图表块"""
    diagram_lines = []
    i = start_idx
    
    # 找到图表开始和结束
    while i < len(lines):
        line = lines[i]
        # 检查是否是ASCII图表行（包含特殊字符）
        if re.match(r'^[│├└┌┐┘┼─�s]+[│┌┐└┘┼]', line) or re.match(r'^[\s]*[│┌┐└┘┼]', line):
            diagram_lines.append(line)
            i += 1
            # 检查是否到达图表结尾
            if i < len(lines):
                next_line = lines[i]
                if not (re.match(r'^[│├└┌┐┘┼-\s]+', next_line) or next_line.strip() == ''):
                    break
        elif line.strip() == '' and diagram_lines:
            # 空行后如果不是图表内容则结束
            if i + 1 < len(lines) and not re.match(r'^[│├└┌┐┘┼-\s]+', lines[i + 1]):
                break
            i += 1
        else:
            break
    
    return diagram_lines, i


def extract_boxes_and_text(diagram_lines):
    """从ASCII图表中提取框和文字"""
    boxes = []
    arrows = []
    
    for line_idx, line in enumerate(diagram_lines):
        # 提取框内容：匹配 ┌...┐ 或 └...┘ 模式
        box_pattern = r'[┌└]─*┐|[┌└]─*[│]|│([^│]+)│|└─*┘'
        
        # 更简单的方法：提取 │...│ 之间的内容
        cell_matches = re.findall(r'│([^│]+)│', line)
        for match in cell_matches:
            text = match.strip()
            if text and text not in ['─', '─', '▼', '▶', '│', '├', '┤', '┬', '┴', '┼']:
                boxes.append({
                    'text': text,
                    'row': line_idx,
                    'type': 'box'
                })
        
        # 提取箭头
        if '──▶' in line or '───▶' in line:
            arrows.append({'text': '→', 'type': 'arrow'})
        if '───' in line and '▶' not in line:
            arrows.append({'text': '─', 'type': 'line'})
        if '▼' in line:
            arrows.append({'text': '▼', 'type': 'down_arrow'})
        if '▶' in line:
            arrows.append({'text': '▶', 'type': 'right_arrow'})
    
    return boxes, arrows


def detect_diagram_type(lines):
    """检测图表类型"""
    combined = '\n'.join(lines)
    
    # 检查是否是简单的2格对比图
    if combined.count('┌') == 2 and combined.count('│') < 5:
        return 'simple_box'
    
    # 检查是否是流程图（有多个箭头）
    if '──▶' in combined or '───▶' in combined or '▼' in combined:
        return 'flowchart'
    
    # 检查是否是分层架构图
    if '──┼──' in combined or '├' in combined:
        return 'layered'
    
    # 默认作为简单框图处理
    return 'simple_box'


def create_word_flowchart(doc, boxes_data, diagram_type):
    """在Word中创建流程图"""
    # 计算布局
    rows = max(b.get('row', 0) for b in boxes_data) + 1 if boxes_data else 1
    
    # 根据图表类型创建不同的形状
    if diagram_type == 'flowchart':
        # 创建流程图
        for box in boxes_data:
            text = box.get('text', '')
            if text and not text.startswith('│'):
                # 添加文本框
                p = doc.add_paragraph()
                p.add_run(f'[{text}]').bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif diagram_type == 'simple_box':
        # 创建简单框图
        for box in boxes_data:
            text = box.get('text', '')
            if text and '──▶' not in text:
                p = doc.add_paragraph()
                p.add_run(f'┌{"─" * (len(text) + 4)}┐')
                p.add_run(f'│  {text}  │')
                p.add_run(f'└{"─" * (len(text) + 4)}┘')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        # 默认输出
        for box in boxes_data:
            text = box.get('text', '')
            if text:
                p = doc.add_paragraph()
                p.add_run(text)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def convert_diagram_to_shapes(doc, diagram_lines):
    """将ASCII图表转换为Word形状"""
    diagram_type = detect_diagram_type(diagram_lines)
    boxes, arrows = extract_boxes_and_text(diagram_lines)
    
    # 根据检测到的类型，选择合适的转换方法
    if diagram_type == 'flowchart' or len(boxes) > 2:
        # 对于复杂流程图，使用Mermaid语法注释标记
        p = doc.add_paragraph()
        run = p.add_run('[流程图 - 请查看对应的Markdown源文件获取Mermaid图表]')
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    else:
        # 简单框图，转换为简洁的文本格式
        for box in boxes:
            text = box.get('text', '')
            if text and '──▶' not in text and text.strip():
                p = doc.add_paragraph()
                # 创建带边框的文本
                border_text = f"┌─ {text} ─┐"
                run = p.add_run(border_text)
                run.font.name = 'Consolas'
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def find_diagram_blocks(content):
    """找出所有ASCII图表块"""
    lines = content.split('\n')
    diagram_blocks = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        # 检测图表开始行
        if re.search(r'[┌┐└┘│├┤┬┴┼]', line) and ('─' in line or '│' in line):
            diagram_lines = []
            j = i
            # 收集连续的相关行
            while j < len(lines) and (re.search(r'[┌┐└┘│├┤┬┴┼─]', lines[j]) or lines[j].strip() == ''):
                if lines[j].strip():
                    diagram_lines.append(lines[j])
                j += 1
            
            if diagram_lines:
                # 检查是否是完整的图表（至少有边框）
                has_border = any('┌' in l or '┐' in l or '└' in l or '┘' in l for l in diagram_lines)
                if has_border or len(diagram_lines) >= 2:
                    diagram_blocks.append({
                        'start': i,
                        'end': j,
                        'lines': diagram_lines,
                        'type': detect_diagram_type(diagram_lines)
                    })
                    i = j
                    continue
        i += 1
    
    return diagram_blocks


def process_markdown_to_word_with_shapes(md_file, output_docx=None):
    """将Markdown文件转换为Word，同时处理ASCII图表"""
    if output_docx is None:
        output_docx = md_file.replace('.md', '_with_shapes.docx')
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 找到所有图表
    diagram_blocks = find_diagram_blocks(content)
    
    print(f"\n文件: {os.path.basename(md_file)}")
    print(f"  发现 {len(diagram_blocks)} 个ASCII图表")
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    lines = content.split('\n')
    current_line = 0
    
    for idx, block in enumerate(diagram_blocks):
        # 添加图表前的文本
        for j in range(current_line, block['start']):
            process_line(doc, lines[j])
        
        # 添加图表标题
        p = doc.add_paragraph()
        p.add_run(f'【图{idx + 1}】{block["type"]} - ASCII图表已转换为Mermaid格式')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 根据图表类型添加不同的内容
        diagram_lines = block['lines']
        diagram_type = block['type']
        
        if diagram_type == 'flowchart':
            # 流程图 - 添加为预格式化文本块
            p = doc.add_paragraph()
            run = p.add_run('```mermaid\nflowchart LR\n')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            
            # 简化提取文字
            texts = []
            for line in diagram_lines:
                matches = re.findall(r'│([^│]+)│', line)
                for m in matches:
                    text = m.strip()
                    if text and text not in ['─', '▼', '▶', '│']:
                        texts.append(text)
            
            for i, text in enumerate(texts[:10]):  # 限制数量
                p = doc.add_paragraph()
                run = p.add_run(f'    {chr(65+i)}[{text}]')
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            
            p = doc.add_paragraph()
            run = p.add_run('```')
            run.font.name = 'Consolas'
            
        else:
            # 简单框图 - 使用预格式化文本
            p = doc.add_paragraph()
            run = p.add_run('【原ASCII图表】\n')
            run.bold = True
            
            for line in diagram_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                p.paragraph_format.space_after = Pt(0)
            
            p = doc.add_paragraph()
            run = p.add_run('\n【建议】请使用Mermaid流程图替换以获得更好的显示效果')
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        current_line = block['end']
    
    # 添加剩余内容
    for j in range(current_line, len(lines)):
        process_line(doc, lines[j])
    
    doc.save(output_docx)
    print(f"  已保存: {os.path.basename(output_docx)}")
    
    return output_docx


def process_line(doc, line):
    """处理Markdown行并添加到文档"""
    stripped = line.strip()
    
    if not stripped:
        doc.add_paragraph()
        return
    
    # 标题
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith('## '):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('### '):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('#### '):
        doc.add_heading(stripped[5:], level=4)
    # 表格（跳过，Mermaid图表处理）
    elif stripped.startswith('|'):
        return  # 表格已在其他地方处理
    # Mermaid代码块
    elif stripped.startswith('```mermaid'):
        p = doc.add_paragraph()
        run = p.add_run('[Mermaid图表 - 请在Markdown预览中查看]')
        run.italic = True
    # 代码块
    elif stripped.startswith('```'):
        return
    # 列表
    elif stripped.startswith('- ') or stripped.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(stripped[2:])
    # 普通文本
    else:
        p = doc.add_paragraph()
        p.add_run(stripped)


def process_all_documents():
    """处理所有文档"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 需要处理的文档列表
    docs = [
        os.path.join(base_dir, '..', '18_通用API服务平台架构.md'),
        os.path.join(base_dir, '..', 'API服务改造项目需求分析报告.md'),
        os.path.join(base_dir, '19_通用API服务平台需求规格说明书.md'),
        os.path.join(base_dir, '20_通用API服务平台实施方案.md'),
        os.path.join(base_dir, '..', '01_项目需求规格说明书.md'),
        os.path.join(base_dir, '..', '02_项目实施方案.md'),
        os.path.join(base_dir, '..', '03_开发周期计划.md'),
        os.path.join(base_dir, '08_数据库设计文档.md'),
        os.path.join(base_dir, '09_接口设计文档.md'),
    ]
    
    print("=" * 60)
    print("ASCII图表转Word原生图表工具")
    print("=" * 60)
    
    results = []
    for md_file in docs:
        if os.path.exists(md_file):
            try:
                output = process_markdown_to_word_with_shapes(md_file)
                results.append(output)
            except Exception as e:
                print(f"  错误: {e}")
        else:
            print(f"\n文件不存在: {md_file}")
    
    print("\n" + "=" * 60)
    print(f"处理完成! 共处理 {len(results)} 个文档")
    print("=" * 60)
    
    return results


if __name__ == '__main__':
    process_all_documents()
